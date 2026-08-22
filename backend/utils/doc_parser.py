"""
Prepo.ai — Document Parser & Text Extraction Utilities

Supports PDF (.pdf), Word Documents (.docx), and Plain Text (.txt) files.
Includes Layer-1 heuristic guardrails for empty, unreadable, or scanned-only documents.
"""

from __future__ import annotations

import io
import os
import re
from typing import Tuple


def extract_text_from_file(file_bytes: bytes, filename: str) -> str:
    """
    Extract readable text from uploaded file bytes based on file extension.
    Supports PDF (.pdf), Word (.docx), and Text (.txt) files.
    """
    ext = os.path.splitext(filename)[1].lower()

    if ext == ".pdf":
        return _extract_from_pdf(file_bytes)
    elif ext in (".docx", ".doc"):
        return _extract_from_docx(file_bytes)
    elif ext in (".txt", ".md", ".csv"):
        return _extract_from_txt(file_bytes)
    else:
        raise ValueError(f"Unsupported file format '{ext}'. Please upload a PDF (.pdf), Word document (.docx), or Text file (.txt).")


def _extract_from_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF bytes using pypdf."""
    try:
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(file_bytes))
        pages_text = []
        for i, page in enumerate(reader.pages):
            text = page.extract_text() or ""
            if text.strip():
                pages_text.append(text.strip())
        return "\n\n".join(pages_text)
    except Exception as e:
        print(f"Error reading PDF with pypdf: {e}")
        raise ValueError(f"Could not read PDF file: {str(e)}. The file may be corrupt or encrypted.")


def _extract_from_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX bytes using python-docx."""
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        # Also extract table cells text
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        return "\n\n".join(paragraphs)
    except Exception as e:
        print(f"Error reading DOCX: {e}")
        raise ValueError(f"Could not read Word document: {str(e)}. Please upload a valid .docx file.")


def _extract_from_txt(file_bytes: bytes) -> str:
    """Extract text from plain text bytes."""
    for encoding in ("utf-8", "utf-8-sig", "latin-1", "cp1252"):
        try:
            return file_bytes.decode(encoding)
        except UnicodeDecodeError:
            continue
    return file_bytes.decode("utf-8", errors="replace")


def validate_document_text(text: str) -> Tuple[bool, str]:
    """
    Layer-1 Fast Heuristic Guardrail.
    Validates that extracted text is non-empty, contains sufficient words,
    and consists of readable educational text rather than blank/scanned pages or gibberish.
    
    Returns (is_valid: bool, error_message: str)
    """
    clean_text = text.strip()

    if not clean_text:
        return False, "The uploaded document is completely empty or contains only scanned images without readable text. Please upload a document with clear text."

    # Word count check
    words = re.findall(r"\b\w+\b", clean_text)
    word_count = len(words)

    if word_count < 40:
        return False, f"The uploaded document contains insufficient text (only {word_count} words found, minimum 40 words required). Please upload a more comprehensive study document or notes."

    # Alphanumeric ratio check (ensures not binary garbage or repeated punctuation)
    alpha_chars = sum(1 for c in clean_text if c.isalnum())
    total_chars = len(clean_text)

    if total_chars > 0 and (alpha_chars / total_chars) < 0.35:
        return False, "The uploaded document contains mostly unreadable symbols or corrupt text. Please upload a clear text document."

    return True, ""


def truncate_document_text(text: str, max_chars: int = 60000) -> str:
    """
    Safely truncate text if it exceeds Groq LLM practical prompt size.
    60,000 characters is ~15,000 words (about 30-40 pages), easily fitting in Llama 3.3 70B.
    """
    if len(text) <= max_chars:
        return text
    return text[:max_chars] + "\n\n... [Document truncated for optimal quiz generation]"
